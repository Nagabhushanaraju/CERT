#!/usr/bin/env python3
"""Reprocess organized dump: convert all files to images, separate resume, dedupe, classify, and add short descriptions for certificates.

Run from workspace root. Example:
  python3 reprocess_and_sort.py \
    --source /workspaces/CERT/dump-/drive-download-20260516T185024Z-3-001_unzipped \
    --out /workspaces/CERT/dump-/final_sorted

This script:
- Finds resume PDFs (filename contains 'resume' or 'cv'), converts them to images and saves under `resume/`.
- Converts PDFs, images, and DOCX into PNG images.
- Deduplicates images by SHA256 of normalized PNG bytes.
- Classifies images into categories using filename and PDF-extracted text (PyMuPDF).
- For items classified as `certificates`, writes a short description (2-5 words) into `certificates/descriptions.json`.
"""
from __future__ import annotations

import argparse
import shutil
import os
from pathlib import Path
import hashlib
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import pytesseract

try:
    import fitz
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None


def sha256_bytes(b: bytes) -> str:
    import hashlib
    return hashlib.sha256(b).hexdigest()


def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)


def pdf_to_images(path: Path, dpi=200) -> list[bytes]:
    imgs = []
    if fitz is None:
        return imgs
    doc = fitz.open(str(path))
    for page in doc:
        pix = page.get_pixmap(dpi=dpi)
        imgs.append(pix.tobytes('png'))
    return imgs


def image_file_to_png_bytes(path: Path) -> bytes:
    img = Image.open(path)
    # autocrop small black borders
    gray = img.convert('L')
    bw = gray.point(lambda p: 255 if p > 10 else 0, '1')
    bbox = bw.getbbox()
    if bbox:
        img = img.crop(bbox)
    # standardize width
    maxw = 1200
    if img.width > maxw:
        h = int(maxw * img.height / img.width)
        img = img.resize((maxw, h), Image.LANCZOS)
    buf = BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def docx_to_image_bytes(path: Path) -> list[bytes]:
    texts = []
    if Document is not None:
        try:
            doc = Document(path)
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
        except Exception:
            texts = []
    if not texts:
        texts = [path.stem]
    # render each paragraph block into an image (or single image)
    full = '\n'.join(texts[:40])
    # create white image and write text
    font = None
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None
    img = Image.new('RGB', (1000, 1200), 'white')
    d = ImageDraw.Draw(img)
    margin = 20
    d.multiline_text((margin, margin), full, fill='black', font=font)
    buf = BytesIO()
    img.save(buf, format='PNG')
    return [buf.getvalue()]


def extract_pdf_text(path: Path) -> str:
    if fitz is None:
        return ''
    try:
        doc = fitz.open(str(path))
        txt = []
        for p in doc:
            txt.append(p.get_text() or '')
        return '\n'.join(txt)
    except Exception:
        return ''


def classify_text_and_name(text: str, name: str) -> str:
    txt = (text or '') + '\n' + name
    t = txt.lower()
    if any(k in t for k in ['certificate', 'certified', 'certificate of']):
        return 'certificates'
    if any(k in t for k in ['publication', 'paper', 'journal', 'proceeding']):
        return 'publications'
    if any(k in t for k in ['intern', 'internship']):
        return 'internships'
    if any(k in t for k in ['award', 'achievement', 'winner']):
        return 'awards'
    if any(k in t for k in ['resume', 'curriculum vitae', 'cv']):
        return 'resume'
    return 'others'


def classify_ocr_text(text: str, name: str) -> str:
    t = ((text or '') + '\n' + name).lower()
    if any(k in t for k in ['resume', 'curriculum vitae', 'cv']):
        return 'resume'
    if any(k in t for k in [
        'certificate', 'certified', 'certificate of', 'certifcate', 'study certificate',
        'study cert', 'this is to certify', 'this certificate', 'to certify that',
        'has successfully completed', 'course completion', 'workshop completion',
        'certificate number', 'certificate for', 'of participation', 'of achievement',
        'completed on', 'completed the', 'completed various courses', 'attests to', 'certify that',
    ]):
        return 'certificates'
    if any(k in t for k in ['publication', 'paper', 'journal', 'proceeding']):
        return 'publications'
    if any(k in t for k in ['intern', 'internship']):
        return 'internships'
    if any(k in t for k in ['award', 'achievement', 'winner']):
        return 'awards'
    if any(k in t for k in ['hackathon', 'workshop', 'participation', 'course', 'learning', 'completion']):
        return 'certificates'
    return 'others'


def short_description_for_certificate(text: str, name: str) -> str:
    # Try to locate a meaningful short phrase from text or name
    src = (text or '').strip()
    if src:
        lines = [l.strip() for l in src.splitlines() if l.strip()]
        for line in lines:
            if 'certificate' in line.lower():
                parts = line.split()
                # take up to 5 words after 'certificate' or first 5 words
                if 'certificate' in parts[0].lower():
                    cand = ' '.join(parts[1:6])
                    if cand:
                        return ' '.join(cand.split()[:5])
                # otherwise return first 3 words of line
                return ' '.join(parts[:5])
        # fallback: first 3 words of first line
        parts = lines[0].split()
        return ' '.join(parts[:5])
    # fallback to filename tokens without extension or 'certificate'
    name = name.replace('_', ' ').replace('-', ' ')
    tokens = [w for w in name.split() if 'cert' not in w.lower()]
    if tokens:
        return ' '.join(tokens[:5])
    return 'Certificate'


def sanitize_description(text: str) -> str:
    words = []
    for raw in text.replace('|', ' ').replace('/', ' ').split():
        cleaned = ''.join(ch for ch in raw if ch.isalnum())
        if cleaned:
            words.append(cleaned)
        if len(words) == 5:
            break
    return ' '.join(words) if words else 'Certificate'


def ocr_png_bytes(image_bytes: bytes) -> str:
    try:
        img = Image.open(BytesIO(image_bytes))
        return pytesseract.image_to_string(img, lang='eng')
    except Exception:
        return ''


def main(source: Path, out: Path):
    ensure_dir(out)
    seen_hashes: set[str] = set()
    cert_descriptions = {}
    for root, dirs, files in os.walk(source):
        for f in files:
            p = Path(root) / f
            suf = p.suffix.lower()
            name = p.name
            # detect resume by name
            is_resume = any(k in name.lower() for k in ['resume', 'cv', 'curriculum'])
            if is_resume and suf == '.pdf':
                # put all resume pages into resume/images
                imgs = pdf_to_images(p)
                for i, b in enumerate(imgs, start=1):
                    h = sha256_bytes(b)
                    if h in seen_hashes:
                        continue
                    seen_hashes.add(h)
                    ocr_text = ocr_png_bytes(b)
                    outp = out / 'resume' / 'images'
                    ensure_dir(outp)
                    (outp / f"{p.stem}_page{i}.png").write_bytes(b)
                continue
            # for everything else, convert to images
            img_bytes_list = []
            if suf == '.pdf':
                text = extract_pdf_text(p)
                img_bytes_list = pdf_to_images(p)
                category = classify_text_and_name(text, name)
            elif suf in ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff', '.gif']:
                try:
                    b = image_file_to_png_bytes(p)
                    img_bytes_list = [b]
                except Exception as e:
                    print('image convert failed', p, e)
                    continue
                category = classify_text_and_name('', name)
            elif suf == '.docx':
                # extract text & render
                text = ''
                if Document is not None:
                    try:
                        doc = Document(p)
                        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    except Exception:
                        text = ''
                img_bytes_list = docx_to_image_bytes(p)
                category = classify_text_and_name(text, name)
            else:
                # copy unknown to others/raw
                target = out / 'others' / 'raw'
                ensure_dir(target)
                h = sha256_bytes(p.read_bytes())
                if h in seen_hashes:
                    continue
                seen_hashes.add(h)
                shutil.copy2(p, target / p.name)
                continue

            for b in img_bytes_list:
                h = sha256_bytes(b)
                if h in seen_hashes:
                    continue
                seen_hashes.add(h)
                ocr_text = ocr_png_bytes(b)
                if ocr_text:
                    category = classify_ocr_text(ocr_text, name)
                elif suf == '.pdf':
                    category = classify_text_and_name(text if 'text' in locals() else '', name)
                outp = out / category / 'images'
                ensure_dir(outp)
                # create filename
                base = p.stem
                # ensure unique name
                idx = 1
                candidate = outp / f"{base}.png"
                while candidate.exists():
                    candidate = outp / f"{base}_{idx}.png"
                    idx += 1
                candidate.write_bytes(b)
                # if certificate, add short description
                if category == 'certificates':
                    desc_source = ocr_text if ocr_text else (text if 'text' in locals() else '')
                    desc = short_description_for_certificate(desc_source, name)
                    cert_descriptions[candidate.name] = sanitize_description(desc)

    # write certificates descriptions
    if cert_descriptions:
        import json
        dest = out / 'certificates'
        ensure_dir(dest)
        with (dest / 'descriptions.json').open('w', encoding='utf8') as fh:
            json.dump(cert_descriptions, fh, indent=2, ensure_ascii=False)

    print('Reprocessing complete. Output at', out)


if __name__ == '__main__':
    p = argparse.ArgumentParser()
    p.add_argument('--source', required=False, default=str(Path.cwd() / 'dump-'), help='unzipped source folder')
    p.add_argument('--out', required=False, default=str(Path.cwd() / 'dump-/final_sorted'), help='output folder')
    args = p.parse_args()
    main(Path(args.source), Path(args.out))
